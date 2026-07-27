"""Build blinded PERSONA-MH human-annotation CSVs and DOCX instructions.

Run from the repository root:
    python annotation_materials/build_annotation_packet.py

The generated rating columns are intentionally blank. Existing automated
PERSONA scores and HuMT values are never included in annotator-facing files.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from persona_annotation.config import load_config
from persona_annotation.loaders import load_all_sources

OUTPUT_DIR = Path(__file__).resolve().parent
PHASE1_SEED = 20260727
PHASE2_SEED = 20260728
PROTOCOL_ID = "persona_mh_human_v2"
RUBRIC_VERSION = "2.0"


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize_stimulus(value: str) -> str:
    """Normalize line endings and trailing whitespace without rewriting text."""

    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


def _clean_markdown(text: str) -> str:
    """Remove lightweight inline Markdown for DOCX rendering."""

    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def markdown_to_docx(source: Path, target: Path) -> None:
    """Render the protocol's simple headings, lists, and tables to DOCX."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 12)):
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)

    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    first_heading = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            separator = lines[i + 1].strip()
            if separator.startswith("|") and "---" in separator:
                headers = [_clean_markdown(c.strip()) for c in stripped.strip("|").split("|")]
                rows: list[list[str]] = []
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    rows.append(
                        [_clean_markdown(c.strip()) for c in lines[i].strip().strip("|").split("|")]
                    )
                    i += 1
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for j, value in enumerate(headers):
                    table.rows[0].cells[j].text = value
                    for run in table.rows[0].cells[j].paragraphs[0].runs:
                        run.bold = True
                for row in rows:
                    cells = table.add_row().cells
                    for j, value in enumerate(row[: len(cells)]):
                        cells[j].text = value
                continue

        if stripped.startswith("# "):
            if not first_heading:
                document.add_page_break()
            document.add_heading(_clean_markdown(stripped[2:]), level=0)
            first_heading = False
        elif stripped.startswith("## "):
            document.add_heading(_clean_markdown(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            document.add_heading(_clean_markdown(stripped[4:]), level=2)
        elif re.match(r"^[-*] ", stripped):
            document.add_paragraph(_clean_markdown(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            value = re.sub(r"^\d+\.\s+", "", stripped)
            document.add_paragraph(_clean_markdown(value), style="List Number")
        elif stripped == "---":
            document.add_paragraph().add_run().add_break(WD_BREAK.LINE)
        else:
            document.add_paragraph(_clean_markdown(stripped))
        i += 1

    core = document.core_properties
    core.title = source.stem.replace("_", " ")
    core.subject = "PERSONA-MH human annotation"
    core.comments = f"Generated from {source.name}; protocol {PROTOCOL_ID}"
    document.save(target)


def build_rows() -> pd.DataFrame:
    """Load configured response sources and assign stable blinded item IDs."""

    cfg = load_config(REPO_ROOT / "persona_annotation" / "config.yaml")
    source_rows = load_all_sources(cfg.sources)
    records = []
    for row in source_rows:
        prompt = _normalize_stimulus(row.prompt)
        response = _normalize_stimulus(row.response)
        digest = hashlib.sha256(
            (
                f"{PROTOCOL_ID}\0{row.source_id}\0{row.prompt_id}\0"
                f"{row.model}\0{prompt}\0{response}"
            ).encode("utf-8")
        ).hexdigest()
        records.append(
            {
                "model": row.model,
                "prompt_id": row.prompt_id,
                "source_id": row.source_id,
                "source_set": row.source_set,
                "topic": row.topic,
                "failure_mode": row.failure_mode or "",
                "prompt": prompt,
                "response": response,
                "response_file": str(Path(row.response_file).relative_to(REPO_ROOT)),
                "source_row_index": row.row_index,
                "content_sha256": digest,
            }
        )

    frame = pd.DataFrame(records)
    if len(frame) != 660:
        raise ValueError(f"Expected 660 configured responses, found {len(frame)}")
    if frame["response"].isna().any() or (frame["response"].str.strip() == "").any():
        raise ValueError("Every annotation item must have a non-empty response")
    if frame["prompt"].isna().any() or (frame["prompt"].str.strip() == "").any():
        raise ValueError("Every annotation item must have a non-empty prompt")
    if frame["content_sha256"].duplicated().any():
        raise ValueError("Duplicate prompt/model/response annotation item detected")

    # IDs are assigned after a fixed shuffle so their sequence does not reveal
    # source-file or model ordering.
    frame = frame.sample(frac=1, random_state=PHASE1_SEED).reset_index(drop=True)
    frame.insert(0, "annotation_item_id", [f"PMH-{i:04d}" for i in range(1, len(frame) + 1)])
    frame.insert(1, "protocol_id", PROTOCOL_ID)
    return frame


def write_csvs(frame: pd.DataFrame) -> None:
    """Write two blinded rating sheets and a private re-identification key."""

    phase1 = frame[["annotation_item_id", "prompt", "response"]].copy()
    phase1.insert(1, "protocol_id", PROTOCOL_ID)
    phase1.insert(2, "rubric_version", RUBRIC_VERSION)
    phase1.insert(3, "presentation_order", range(1, len(phase1) + 1))
    phase1.insert(4, "annotator_id", "")
    phase1["OA_score"] = ""
    phase1["OA_reason"] = ""
    phase1["OA_review_flag"] = ""
    phase1["annotator_notes"] = ""
    phase1.to_csv(
        OUTPUT_DIR / "persona_mh_phase1_oa_annotation.csv",
        index=False,
        encoding="utf-8-sig",
    )

    # Use a separate deterministic order to reduce recall/order carryover.
    phase2 = frame[["annotation_item_id", "prompt", "response"]].sample(
        frac=1, random_state=PHASE2_SEED
    ).reset_index(drop=True)
    phase2.insert(1, "protocol_id", PROTOCOL_ID)
    phase2.insert(2, "rubric_version", RUBRIC_VERSION)
    phase2.insert(3, "presentation_order", range(1, len(phase2) + 1))
    phase2.insert(4, "annotator_id", "")
    for column in (
        "scenario_type",
        "E_score",
        "E_reason",
        "E_evidence",
        "D_score",
        "D_reason",
        "D_evidence",
        "F_score",
        "F_reason",
        "F_evidence",
        "review_flag",
        "annotator_notes",
    ):
        phase2[column] = ""
    phase2.to_csv(
        OUTPUT_DIR / "persona_mh_phase2_edf_annotation.csv",
        index=False,
        encoding="utf-8-sig",
    )

    phase1_order = dict(zip(phase1["annotation_item_id"], phase1["presentation_order"]))
    phase2_order = dict(zip(phase2["annotation_item_id"], phase2["presentation_order"]))
    frame = frame.copy()
    frame["phase1_template_order"] = frame["annotation_item_id"].map(phase1_order)
    frame["phase2_template_order"] = frame["annotation_item_id"].map(phase2_order)
    key_columns = [
        "annotation_item_id",
        "protocol_id",
        "model",
        "prompt_id",
        "source_id",
        "source_set",
        "topic",
        "failure_mode",
        "response_file",
        "source_row_index",
        "content_sha256",
        "phase1_template_order",
        "phase2_template_order",
    ]
    frame[key_columns].to_csv(
        OUTPUT_DIR / "persona_mh_item_key_DO_NOT_SHARE_WITH_ANNOTATORS.csv",
        index=False,
        encoding="utf-8-sig",
    )

    source_paths = sorted({REPO_ROOT / value for value in frame["response_file"]})
    manifest = {
        "protocol_id": PROTOCOL_ID,
        "rubric_version": RUBRIC_VERSION,
        "n_items": len(frame),
        "phase1_seed": PHASE1_SEED,
        "phase2_seed": PHASE2_SEED,
        "config_file": "persona_annotation/config.yaml",
        "config_sha256": _sha256_file(REPO_ROOT / "persona_annotation" / "config.yaml"),
        "source_files": [
            {
                "path": str(path.relative_to(REPO_ROOT)),
                "sha256": _sha256_file(path),
            }
            for path in source_paths
        ],
        "annotator_files_exclude": [
            "model",
            "humt_score",
            "automated PERSONA scores",
        ],
    }
    (OUTPUT_DIR / "persona_mh_build_manifest.json").write_text(
        json.dumps(manifest, indent=2) + "\n", encoding="utf-8"
    )


def make_annotator_copy(
    *,
    phase: str,
    annotator_id: str,
    order_seed: int,
    output_path: Path,
) -> None:
    """Create one assigned, independently randomized annotator CSV."""

    if not annotator_id.strip():
        raise ValueError("annotator_id cannot be blank")
    template_name = (
        "persona_mh_phase1_oa_annotation.csv"
        if phase == "phase1"
        else "persona_mh_phase2_edf_annotation.csv"
    )
    template_path = OUTPUT_DIR / template_name
    if not template_path.exists():
        write_csvs(build_rows())
    frame = pd.read_csv(template_path, keep_default_na=False)
    frame = frame.sample(frac=1, random_state=order_seed).reset_index(drop=True)
    frame["presentation_order"] = range(1, len(frame) + 1)
    frame["annotator_id"] = annotator_id.strip()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    frame.to_csv(output_path, index=False, encoding="utf-8-sig")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--make-copy", choices=("phase1", "phase2"))
    parser.add_argument("--annotator-id")
    parser.add_argument("--order-seed", type=int)
    parser.add_argument("--output", type=Path)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    if args.make_copy:
        if args.annotator_id is None or args.order_seed is None or args.output is None:
            raise SystemExit(
                "--make-copy requires --annotator-id, --order-seed, and --output"
            )
        make_annotator_copy(
            phase=args.make_copy,
            annotator_id=args.annotator_id,
            order_seed=args.order_seed,
            output_path=args.output,
        )
        print(f"Built {args.make_copy} copy for {args.annotator_id} → {args.output}")
        return

    frame = build_rows()
    write_csvs(frame)
    for markdown_path in (
        OUTPUT_DIR / "PERSONA_MH_Phase1_OA_Instructions.md",
        OUTPUT_DIR / "PERSONA_MH_Human_Annotation_Protocol_v2.md",
    ):
        markdown_to_docx(markdown_path, markdown_path.with_suffix(".docx"))
    print(f"Built annotation packet for {len(frame)} responses in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
